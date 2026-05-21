"""Research-document knowledge pipeline — extraction + tagging.

Walks the ``research/`` folder, extracts text from PDFs (pypdf, with a
pdfplumber fallback for stubborn pages) and Word docs (python-docx), and for
each NEW document (deduped by SHA-256 against the ``documents`` table):

- writes framework study-notes to ``docs/playbooks/<slug>.md``
- writes ``themes`` + ``theme_observations`` rows (auto-tag + summary)

Embeddings / RAG are intentionally deferred to the next layer. The pipeline is
idempotent: a document already in ``documents`` (same sha256) is skipped, so
re-runs are safe (CLAUDE.md rule 5 spirit).

Documents are flagged ``source="internal"`` — research material may be
proprietary / licensed (CLAUDE.md R5); never export the extracted text.

Manual run (Ollama must be running locally — it lives on the laptop):
    python -m trading_intel.memory.pdf_pipeline
    python -m trading_intel.memory.pdf_pipeline --limit 3
    python -m trading_intel.memory.pdf_pipeline --research-dir research --model qwen2.5:3b
"""

from __future__ import annotations

import argparse
import hashlib
import re
import warnings
from datetime import date
from pathlib import Path

import structlog
from sqlalchemy import select
from sqlalchemy.orm import Session

from trading_intel.config import Settings, get_settings
from trading_intel.errors import TradingIntelError
from trading_intel.memory.models import Document, Theme, ThemeObservation
from trading_intel.synthesis.llm import LLMProvider
from trading_intel.synthesis.tagging import (
    MAX_CHARS,
    DocTags,
    extract_framework,
    tag_document,
)

log = structlog.get_logger(__name__)

PDF_EXTS = {".pdf"}
DOCX_EXTS = {".docx"}
SUPPORTED_EXTS = PDF_EXTS | DOCX_EXTS
_SOURCE = "internal"
_MIN_USABLE_CHARS = 100


def discover_documents(research_dir: Path) -> list[Path]:
    """Return supported research files (PDF/docx), sorted, ignoring images etc."""
    if not research_dir.is_dir():
        raise TradingIntelError(f"Research dir not found: {research_dir}")
    return sorted(
        p for p in research_dir.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED_EXTS
    )


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for block in iter(lambda: fh.read(1 << 20), b""):
            digest.update(block)
    return digest.hexdigest()


def slugify(name: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")
    return slug or "doc"


def extract_pdf(path: Path) -> tuple[str, int]:
    """Extract text + page count. Falls back to pdfplumber for empty/scanned PDFs."""
    from pypdf import PdfReader

    with warnings.catch_warnings():
        warnings.simplefilter("ignore")
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        parts: list[str] = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception as exc:  # one bad page must not kill the document
                log.debug("pdf.page_failed", path=path.name, error=str(exc))
                parts.append("")
        text = "\n".join(parts)

    if len(text.strip()) < 50:  # likely scanned / image-only -> try pdfplumber
        try:
            import pdfplumber

            with pdfplumber.open(str(path)) as pdf:
                text = "\n".join((p.extract_text() or "") for p in pdf.pages)
                page_count = len(pdf.pages)
        except Exception as exc:
            log.warning("pdf.plumber_failed", path=path.name, error=str(exc))
    return text, page_count


def extract_docx(path: Path) -> tuple[str, int]:
    """Extract docx paragraph + table text. Page count is not available -> 0."""
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts), 0


def extract_text(path: Path) -> tuple[str, int]:
    ext = path.suffix.lower()
    if ext in PDF_EXTS:
        return extract_pdf(path)
    if ext in DOCX_EXTS:
        return extract_docx(path)
    raise TradingIntelError(f"Unsupported document type: {path.name}")


def _get_or_create_theme(session: Session, name: str, scope: str) -> Theme:
    existing = session.execute(select(Theme).where(Theme.name == name)).scalar_one_or_none()
    if existing is not None:
        return existing
    theme = Theme(name=name, scope=scope)
    session.add(theme)
    session.flush()  # assign theme.id for the FK
    return theme


def _write_playbook(
    playbook_dir: Path, slug: str, title: str, source_name: str, markdown: str, *, truncated: bool
) -> Path:
    playbook_dir.mkdir(parents=True, exist_ok=True)
    out = playbook_dir / f"{slug}.md"
    note = (
        " Generated from the opening section only (document exceeds the model's " "context window)."
        if truncated
        else ""
    )
    header = (
        f"# {title}\n\n"
        f"*Auto-generated framework notes from `research/{source_name}`. Review before "
        f"relying on.{note}*\n\n"
    )
    out.write_text(header + markdown + "\n", encoding="utf-8")
    return out


def ingest_document(
    session: Session,
    llm: LLMProvider,
    path: Path,
    *,
    playbook_dir: Path,
    kind: str = "methodology",
    model: str | None = None,
) -> str:
    """Process one document. Returns ``ingested`` | ``skipped`` | ``empty``."""
    sha = sha256_file(path)
    existing = session.execute(select(Document).where(Document.sha256 == sha)).scalar_one_or_none()
    if existing is not None:
        log.info("pipeline.skip_existing", file=path.name)
        return "skipped"

    text, page_count = extract_text(path)
    if len(text.strip()) < _MIN_USABLE_CHARS:
        log.warning("pipeline.empty_text", file=path.name, chars=len(text.strip()))
        return "empty"

    title = path.stem
    doc = Document(
        path=str(path),
        source=_SOURCE,
        type=path.suffix.lower().lstrip("."),
        kind=kind,
        sha256=sha,
        page_count=page_count or None,
    )
    session.add(doc)
    session.flush()  # assign doc.id

    framework_md = extract_framework(llm, title, text, model=model)
    _write_playbook(
        playbook_dir,
        slugify(title),
        title,
        path.name,
        framework_md,
        truncated=len(text) > MAX_CHARS,
    )

    tags: DocTags = tag_document(llm, title, text, model=model)
    obs_date = date.today()
    symbol = tags.symbols[0][:16] if tags.symbols else None
    for tag in tags.themes:
        theme = _get_or_create_theme(session, tag.name, tag.scope)
        session.add(
            ThemeObservation(
                theme_id=theme.id,
                symbol=symbol,
                date=obs_date,
                sentiment=tag.sentiment,
                source_doc_id=doc.id,
                quote_text=tags.summary or None,
                confidence=tag.confidence,
            )
        )

    session.commit()
    log.info(
        "pipeline.ingested",
        file=path.name,
        pages=page_count,
        themes=len(tags.themes),
        symbols=tags.symbols,
    )
    return "ingested"


def run(
    session: Session,
    llm: LLMProvider,
    *,
    research_dir: Path,
    playbook_dir: Path,
    limit: int | None = None,
    kind: str = "methodology",
    model: str | None = None,
) -> dict[str, int]:
    """Ingest every supported document under ``research_dir``."""
    docs = discover_documents(research_dir)
    if limit:
        docs = docs[:limit]
    log.info("pipeline.start", n_docs=len(docs), research_dir=str(research_dir))

    stats = {"ingested": 0, "skipped": 0, "empty": 0, "failed": 0}
    for path in docs:
        try:
            status = ingest_document(
                session, llm, path, playbook_dir=playbook_dir, kind=kind, model=model
            )
            stats[status] += 1
        except (TradingIntelError, OSError, ValueError) as exc:
            session.rollback()
            stats["failed"] += 1
            log.warning("pipeline.failed", file=path.name, error=str(exc))
    log.info("pipeline.done", **stats)
    return stats


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Ingest research docs: extract frameworks + tag themes."
    )
    parser.add_argument("--research-dir", default="research")
    parser.add_argument("--playbook-dir", default="docs/playbooks")
    parser.add_argument("--limit", type=int, default=None)
    parser.add_argument(
        "--kind",
        choices=["methodology", "research"],
        default="methodology",
        help="Knowledge type: methodology (LLM reasoning) or research (company/theme).",
    )
    parser.add_argument(
        "--model", default=None, help="Ollama model override (default: LLM_DAILY_MODEL)"
    )
    args = parser.parse_args()

    from trading_intel.memory.db import make_session_factory
    from trading_intel.synthesis.llm import OllamaProvider

    structlog.configure(
        processors=[
            structlog.contextvars.merge_contextvars,
            structlog.processors.add_log_level,
            structlog.processors.TimeStamper(fmt="iso"),
            structlog.processors.JSONRenderer(),
        ]
    )

    settings: Settings = get_settings()
    llm = OllamaProvider(settings)
    session_factory = make_session_factory(settings)
    with session_factory() as session:
        stats = run(
            session,
            llm,
            research_dir=Path(args.research_dir),
            playbook_dir=Path(args.playbook_dir),
            limit=args.limit,
            kind=args.kind,
            model=args.model,
        )
    print(f"Done: {stats}")


if __name__ == "__main__":
    main()
